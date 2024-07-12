from openai import OpenAI
import docx
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

# Set up OpenAI API key
api_key = ''
client = OpenAI(api_key=api_key)

def generate_sow_content(project_requirements):
    # Use GPT to generate content based on project requirements
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant that generates content for Statement of Work documents."},
            {"role": "user", "content": f"Generate a SOW based on these project requirements: {project_requirements}"}
        ]
    )
    return response.choices[0].message.content

def create_sow_document(content):
    # Create a new Word document
    doc = docx.Document()

    # Add a title
    doc.add_heading('Statement of Work', 0)

    # Add content to the document
    sections = content.split('\n\n')
    for section in sections:
        if section.startswith('#'):
            # This is a heading
            level = section.count('#')
            text = section.lstrip('#').strip()
            doc.add_heading(text, level)
        else:
            # This is regular text
            doc.add_paragraph(section)

    # Save the document
    doc.save('oxford copy.docx')

# Main execution
project_requirements = input("Enter your project requirements: ")
generated_content = generate_sow_content(project_requirements)
create_sow_document(generated_content)
print("SOW document has been generated and saved as 'generated_sow.docx'")