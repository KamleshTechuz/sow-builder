from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

def create_sow_index(project_name, version):
    # Create a new Word document
    doc = Document()

    # Add title
    doc.add_heading('Business Requirements Specifications', 0)
    
    # Add project name and version
    doc.add_paragraph(f"{project_name}")
    doc.add_paragraph(f"Version {version}")

    # Add table of contents
    toc = [
        ("1. Introduction", 1),
        ("2. module_1", 1),
        ("a. sub_module_1", 2),
        ("1. topic_1", 3),
        ("2. topic_2", 3),
        ("3. topic_3", 3),
        ("b. sub_module_2", 2),
        ("1. topic_1", 3),
        ("3. module_2", 1),
        ("a. sub_module_1", 2),
        ("1. topic_1", 3),
        ("2. topic_2", 3),
        ("3. topic_3", 3),
        ("b. sub_module_2", 2),
        ("1. topic_1", 3),
    ]

    for item, level in toc:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(level * 12)  # Indent based on level
        p.add_run(item)

    # Save the document
    doc.save('sow_index.docx')

# Use the function
create_sow_index("Your Project Name", "1.0")
print("SOW index has been generated and saved as 'sow_index.docx'")