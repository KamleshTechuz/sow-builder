from openai import OpenAI
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

# OpenAI API setup
api_key = ''
client = OpenAI(api_key=api_key)

# Function to generate text with GPT-3/4
def generate_text(prompt):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt},
        ],
        max_tokens=1500,
        n=1,
        stop=None,
        temperature=0.7
    )

    print('response.choices[0] >>>>>>>>>>>>>>>>>>>>>>>> ', response.choices[0].message.content)

    return response.choices[0].message.content

# Load the template document
template_path = 'oxford.docx'  # Path to your template DOCX file
doc = Document(template_path)

# Function to add content to existing headers
def add_header_content(section, text):
    header = section.header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Arial'

# Function to add content to existing footers
def add_footer_content(section, text):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Arial'

# Function to add a heading
def add_heading(text, level):
    heading = doc.add_heading(level=level)
    heading_run = heading.add_run(text)
    heading_run.bold = True
    heading_run.font.size = Pt(14)

# Function to add a paragraph
def add_paragraph(text):
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Generate sections based on project details
project_details = {
    "Introduction": "Introduce the project here...",
    "System Actors": ["Therapist", "Super Admin"],
    "Therapist User Journey": "User journey details...",
    "Super Admin Features": "Super Admin features details..."
}

for section, content in project_details.items():
    add_heading(section, level=2)
    if isinstance(content, list):
        for item in content:
            add_heading(item, level=3)
            generated_text = generate_text(f"Describe the role and features of {item}.")
            add_paragraph(generated_text)
    else:
        generated_text = generate_text(f"Describe the {section.lower()}.")
        add_paragraph(generated_text)

# Add header and footer content if needed
for section in doc.sections:
    add_header_content(section, 'Header Content Here')
    add_footer_content(section, 'Footer Content Here')

# Save the document
output_path = 'Generated_SOW_Oxford_Mtrain_Web_App_Development.docx'
doc.save(output_path)
