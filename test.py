from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# OpenAI API setup
api_key = ''
client = OpenAI(api_key=api_key)

# Function to generate text with GPT-3/4
def generate_text(prompt):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt},
        ]
        # max_tokens=1500,
        # n=1,
        # stop=None,
        # temperature=0.7
    )
    print('response.choices[0] >>>>>>>>>>>>>>>>>>>>>>>> ', response.choices[0].message.content)
    return response.choices[0].message.content

# Create a new Document
doc = Document()

# Function to add a title
def add_title(text):
    title = doc.add_heading(level=1)
    title_run = title.add_run(text)
    title_run.bold = True
    title_run.font.size = Pt(24)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Function to add a heading
def add_heading(text, level):
    heading = doc.add_heading(level=level)
    heading_run = heading.add_run(text)
    heading_run.bold = True
    heading_run.font.size = Pt(14)

# Function to add a paragraph
def add_paragraph(text):
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Add document sections
add_title('Business Requirements Specifications')
add_paragraph('Oxford Mtrain Web App Development\nVersion 1.0')

# Generate sections based on project details
project_details = {
    "introduction": "Introduce the project here...",
    "system_actors": ["Therapist", "Super Admin"],
    "therapist_user_journey": "User journey details...",
    "super_admin_features": "Super Admin features details..."
}

for section, content in project_details.items():
    add_heading(section.replace('_', ' ').title(), level=2)
    if isinstance(content, list):
        for item in content:
            add_heading(item, level=3)
            generated_text = generate_text(f"Describe the role and features of {item}.")
            add_paragraph(generated_text)
    else:
        generated_text = generate_text(f"Describe the {section.replace('_', ' ')}.")
        add_paragraph(generated_text)

# Save the document
doc.save('SOW_Oxford_Mtrain_Web_App_Development.docx')
